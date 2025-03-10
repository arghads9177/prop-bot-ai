import streamlit as st
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from docx import Document
from docx.shared import Pt
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import re
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

def format_bold(text):
    """Replaces text surrounded by ** with bold formatting tags for PDFs and DOCX."""
    return re.sub(r"\\*\\*(.*?)\\*\\*", r"<b>\1</b>", text)  # For PDFs

def format_bold_docx(para, text):
    """Applies bold formatting to text between ** for DOCX."""
    parts = re.split(r"(\\*\\*.*?\\*\\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

def generate_proposal(proposal_type, about_you, job_requirements):
    """Generates a proposal using OpenAI's GPT-4o model via LangChain."""
    if not about_you.strip():
        return "Please provide details about yourself or your business."
    if not job_requirements.strip():
        return "Please enter the job requirements."
    
    prompt_templates = {
        "Freelancing Proposal": PromptTemplate(
            input_variables=["about_you", "job_requirements"],
            template=(
                "You are an expert proposal writer. Generate a professional freelancing proposal. "
                "If the freelancer's background, skills, or experience are missing, return an appropriate message asking for them. "
                "If job requirements are unclear, request clarification. "
                "Otherwise, structure the proposal with an introduction, freelancer's skills, relevant experience, proposed approach, and closing statement.\n\n"
                "Freelancer Details: {about_you}\n"
                "Job Requirements: {job_requirements}"
            ),
        ),
        "Business Proposal": PromptTemplate(
            input_variables=["about_you", "job_requirements"],
            template=(
                "You are an expert business proposal writer. Generate a professional business proposal. "
                "If business details like description, achievements, or experience are missing, return an appropriate message asking for them. "
                "If job requirements are unclear, request clarification. "
                "Otherwise, structure the proposal with an introduction, company background, services offered, proposed approach, and closing statement.\n\n"
                "Business Details: {about_you}\n"
                "Job Requirements: {job_requirements}"
            ),
        ),
    }

    llm = ChatOpenAI(model_name="gpt-4o")
    prompt = prompt_templates[proposal_type]
    chain = prompt | llm
    response = chain.invoke({"about_you": about_you, "job_requirements": job_requirements})
    return response.content.strip()

# Function to save proposal as DOCX
def get_docx(proposal):
    doc = Document()
    doc.add_heading("Proposal", level=1)
    paragraphs = proposal.split("\n")
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph()
            format_bold_docx(para, para_text)
            para.runs[0].font.size = Pt(12)  # Set font size for clarity
    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream

# Function to save proposal as PDF (Proper Formatting)
def get_pdf(proposal):
    pdf_stream = BytesIO()
    doc = SimpleDocTemplate(pdf_stream, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    custom_style = ParagraphStyle(name="CustomStyle", fontSize=12, spaceAfter=10)
    heading_style = ParagraphStyle(name="HeadingStyle", fontSize=16, spaceAfter=10, fontName="Helvetica-Bold")
    elements = []
    
    elements.append(Paragraph("Proposal", heading_style))
    elements.append(Spacer(1, 0.2 * inch))
    
    paragraphs = proposal.split("\n")
    for para in paragraphs:
        if para.strip():
            elements.append(Paragraph(format_bold(para), custom_style))
            elements.append(Spacer(1, 0.1 * inch))
    
    doc.build(elements)
    pdf_stream.seek(0)
    return pdf_stream

# Streamlit UI
def main():
    st.set_page_config(page_title="PropBot AI - Proposal Generator", layout="centered")
    proposal_type = st.sidebar.radio("Select Proposal Type:", ["Freelancing Proposal", "Business Proposal"], index=0)
    st.title("🚀 PropBot AI - Proposal Generator")
    st.subheader(proposal_type)
    
    about_you = st.text_area("About You:" if proposal_type == "Freelancing Proposal" else "About Your Business:", "", help="Describe your background, skills, and experience" if proposal_type == "Freelancing Proposal" else "Describe your business, achievements, and experience")
    job_requirements = st.text_area("Job Requirements:", "", help="Paste the job requirements you’re applying for")
    
    if "proposal" not in st.session_state:
        st.session_state.proposal = ""
    
    if st.button("Generate Proposal"):
        with st.spinner("Generating proposal..."):
            st.session_state.proposal = generate_proposal(proposal_type, about_you, job_requirements)
    
    if st.session_state.proposal:
        st.text_area("Generated Proposal:", st.session_state.proposal, height=300)
        
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("Download as DOCX", get_docx(st.session_state.proposal), "proposal.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col2:
            st.download_button("Download as PDF", get_pdf(st.session_state.proposal), "proposal.pdf", "application/pdf")
        
        if st.button("Reset"):
            st.session_state.proposal = ""
            st.rerun()

if __name__ == "__main__":
    main()
